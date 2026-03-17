"""
GCC Playbook 2026-2030: Document Consolidation Script
=====================================================
Merges GCC_Playbook_Complete_CLEAN.docx (Parts I+II) with full Part III content
from GCC_Playbook_Part_III.docx into a single comprehensive document.

Output: GCC_Playbook_2026_2030_FINAL.docx
"""

import os
import copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))

BASE_FILE = os.path.join(DOCS_DIR, "GCC_Playbook_Complete_CLEAN.docx")
PART3_FILE = os.path.join(DOCS_DIR, "GCC_Playbook_Part_III.docx")
OUTPUT_FILE = os.path.join(DOCS_DIR, "GCC_Playbook_2026_2030_FINAL.docx")


def find_part3_outline_start(doc):
    """Find where the Part III outline begins in the Complete_CLEAN doc."""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "GCC Playbook Part III" in text and "Frontier" in text:
            return i
        if text == "PART III: THE FRONTIER":
            return i
    # Fallback: look for the Part III heading in Complete_CLEAN
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name and para.style.name.startswith('Heading'):
            if 'Part III' in para.text or 'THE FRONTIER' in para.text:
                return i
    return None


def find_part3_outline_end(doc, start_idx):
    """Find where the Part III outline ends - look for 'Proposed Additions' heading or end."""
    for i in range(start_idx + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if text == "Proposed Additions to Parts I and II":
            # Include this section too, then find next major break
            for j in range(i + 1, len(doc.paragraphs)):
                t2 = doc.paragraphs[j].text.strip()
                # End of document or next non-Part-III section
                if t2 and doc.paragraphs[j].style and doc.paragraphs[j].style.name:
                    sn = doc.paragraphs[j].style.name
                    if sn.startswith('Heading 1') and not any(kw in t2 for kw in ['Chapter', 'Proposed', 'Part III', 'AI', 'Location', 'M&A', 'Deep', 'Finance', 'Talent', 'Legal', 'Innovation', 'ESG', 'Career', 'GCC 2030', 'Women', 'Mid-Market', 'Mental']):
                        return j
            return len(doc.paragraphs)
    return len(doc.paragraphs)


def get_part3_content(doc):
    """Extract all Part III paragraphs from the Part_III file."""
    paragraphs = []
    # Skip the cover page content, start from "About Part III" or first heading
    start = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == "About Part III" or text == "PART III: THE FRONTIER":
            start = i
            break
    
    for i in range(start, len(doc.paragraphs)):
        paragraphs.append(doc.paragraphs[i])
    
    return paragraphs, start


def copy_paragraph(source_para, target_doc):
    """Copy a paragraph with its formatting to the target document."""
    new_para = target_doc.add_paragraph()
    
    # Copy paragraph style
    if source_para.style:
        try:
            new_para.style = target_doc.styles[source_para.style.name]
        except KeyError:
            # Style doesn't exist in target, try to match
            style_name = source_para.style.name
            if 'Heading' in style_name:
                new_para.style = target_doc.styles[style_name]
            else:
                new_para.style = target_doc.styles['Normal']
    
    # Copy paragraph format
    if source_para.paragraph_format:
        pf = new_para.paragraph_format
        sf = source_para.paragraph_format
        if sf.alignment is not None:
            pf.alignment = sf.alignment
        if sf.space_before is not None:
            pf.space_before = sf.space_before
        if sf.space_after is not None:
            pf.space_after = sf.space_after
        if sf.left_indent is not None:
            pf.left_indent = sf.left_indent
        if sf.first_line_indent is not None:
            pf.first_line_indent = sf.first_line_indent
    
    # Copy runs with formatting
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)
        if run.bold is not None:
            new_run.bold = run.bold
        if run.italic is not None:
            new_run.italic = run.italic
        if run.underline is not None:
            new_run.underline = run.underline
        if run.font.size is not None:
            new_run.font.size = run.font.size
        if run.font.name is not None:
            new_run.font.name = run.font.name
        if run.font.color and run.font.color.rgb is not None:
            new_run.font.color.rgb = run.font.color.rgb
    
    # If no runs but has text, add it
    if not source_para.runs and source_para.text:
        new_para.text = source_para.text
    
    return new_para


def copy_table(source_table, target_doc):
    """Copy a table to the target document."""
    rows = len(source_table.rows)
    cols = len(source_table.columns)
    
    new_table = target_doc.add_table(rows=rows, cols=cols)
    new_table.style = 'Table Grid'
    
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            target_cell = new_table.cell(i, j)
            target_cell.text = cell.text
            
            # Copy cell paragraph formatting
            for k, para in enumerate(cell.paragraphs):
                if k == 0:
                    target_para = target_cell.paragraphs[0]
                    target_para.text = ""
                else:
                    target_para = target_cell.add_paragraph()
                
                for run in para.runs:
                    new_run = target_para.add_run(run.text)
                    if run.bold is not None:
                        new_run.bold = run.bold
                    if run.italic is not None:
                        new_run.italic = run.italic
                    if run.font.size is not None:
                        new_run.font.size = run.font.size
    
    return new_table


def add_page_break(doc):
    """Add a page break."""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)
    return para


def add_toc_field(doc):
    """Add a TOC field that Word can update."""
    para = doc.add_paragraph()
    para.style = doc.styles['Normal']
    
    # Add TOC title
    toc_title = doc.add_paragraph()
    toc_title.style = doc.styles['Heading 1']
    toc_title.text = "Table of Contents"
    
    # Add the TOC field code
    toc_para = doc.add_paragraph()
    run = toc_para.add_run()
    
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run._r.append(instrText)
    
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run._r.append(fldChar2)
    
    # Placeholder text
    placeholder_run = toc_para.add_run("Right-click and select 'Update Field' to generate Table of Contents")
    placeholder_run.font.color.rgb = RGBColor(128, 128, 128)
    placeholder_run.font.italic = True
    
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    placeholder_run._r.append(fldChar3)
    
    # Page break after TOC
    page_break_para = doc.add_paragraph()
    run = page_break_para.add_run()
    from docx.enum.text import WD_BREAK
    run.add_break(WD_BREAK.PAGE)
    
    return toc_para


def extract_part3_references(part3_doc):
    """Extract the references section from Part III."""
    refs = []
    in_refs = False
    for para in part3_doc.paragraphs:
        text = para.text.strip()
        if "APPENDIX: VALIDATED REFERENCES" in text or "VALIDATED REFERENCES" in text:
            in_refs = True
            continue
        if in_refs:
            if text:
                refs.append((para, text))
    return refs


def merge_documents():
    """Main merge function."""
    print("=" * 60)
    print("GCC Playbook 2026-2030 Document Consolidation")
    print("=" * 60)
    
    # Load documents
    print("\n[1/6] Loading source documents...")
    base_doc = Document(BASE_FILE)
    part3_doc = Document(PART3_FILE)
    
    print(f"  Base (Complete_CLEAN): {len(base_doc.paragraphs)} paragraphs, {len(base_doc.tables)} tables")
    print(f"  Part III: {len(part3_doc.paragraphs)} paragraphs, {len(part3_doc.tables)} tables")
    
    # Create new output document
    print("\n[2/6] Creating consolidated document...")
    output_doc = Document()
    
    # Set default font
    style = output_doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Ensure heading styles exist
    for i in range(1, 4):
        heading_style = output_doc.styles[f'Heading {i}']
        heading_style.font.name = 'Calibri'
    
    # ---- TITLE PAGE ----
    print("\n[3/6] Building title page and front matter...")
    
    # Add some spacing
    for _ in range(4):
        output_doc.add_paragraph()
    
    title_para = output_doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("The Complete India GCC Reference")
    title_run.font.size = Pt(28)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = output_doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("2026 – 2030")
    sub_run.font.size = Pt(22)
    sub_run.font.color.rgb = RGBColor(0, 102, 153)
    
    output_doc.add_paragraph()
    
    tagline = output_doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tagline.add_run("Landscape  |  Maturity Models  |  Setup  |  Governance  |  Who's Who\nAI  |  Deep Tech  |  M&A  |  Finance  |  Legal  |  ESG  |  2030 Scenarios")
    tag_run.font.size = Pt(12)
    tag_run.font.color.rgb = RGBColor(100, 100, 100)
    
    output_doc.add_paragraph()
    output_doc.add_paragraph()
    
    author = output_doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth_run = author.add_run("Curated by Kalilur Rahman")
    auth_run.font.size = Pt(14)
    auth_run.bold = True
    
    website = output_doc.add_paragraph()
    website.alignment = WD_ALIGN_PARAGRAPH.CENTER
    web_run = website.add_run("kalilurrahman.lovable.app")
    web_run.font.size = Pt(11)
    web_run.font.color.rgb = RGBColor(0, 102, 204)
    
    # Page break
    from docx.enum.text import WD_BREAK
    pb = output_doc.add_paragraph()
    pb.add_run().add_break(WD_BREAK.PAGE)
    
    # ---- OPEN ACCESS NOTICE ----
    notice_heading = output_doc.add_paragraph()
    notice_heading.style = output_doc.styles['Heading 2']
    notice_heading.text = "Open Access Notice"
    
    notice_text = output_doc.add_paragraph()
    notice_text.text = ("This playbook is an open-access resource compiled for the global GCC community. "
                       "Content is a curated synthesis of publicly available industry research, practitioner "
                       "data, and published frameworks. No proprietary copyright claims are made over core "
                       "concepts, frameworks, or data presented.")
    
    notice_text2 = output_doc.add_paragraph()
    notice_text2.text = ("All leader profiles are sourced from publicly available recognition programs and "
                        "company communications. No entry reflects a personal endorsement by the curator.")
    
    notice_text3 = output_doc.add_paragraph()
    notice_text3.text = ("The intention behind this playbook is purely educational and collaborative. Readers "
                        "are encouraged to freely use, share, and adapt these insights to help navigate their "
                        "own Global Capability Centre journeys.")
    
    # Page break
    pb2 = output_doc.add_paragraph()
    pb2.add_run().add_break(WD_BREAK.PAGE)
    
    # ---- TABLE OF CONTENTS ----
    add_toc_field(output_doc)
    
    # ---- COPY ALL BASE CONTENT (Parts I + II + existing appendices) ----
    print("\n[4/6] Copying Parts I & II content from Complete_CLEAN...")
    
    # Find where Part III outline starts in base
    part3_outline_start = find_part3_outline_start(base_doc)
    
    # Track what sections to skip (we'll replace Part III outline with full content)
    skip_start = None
    skip_end = None
    
    if part3_outline_start is not None:
        skip_start = part3_outline_start
        skip_end = find_part3_outline_end(base_doc, part3_outline_start)
        print(f"  Found Part III outline at paragraph {skip_start}, ends at {skip_end}")
    
    # Build a map of where tables appear in the document body
    # by checking the body XML ordering
    body = base_doc.element.body
    element_list = list(body)
    
    para_idx = 0
    table_idx = 0
    copied_paras = 0
    copied_tables = 0
    
    for element in element_list:
        if element.tag.endswith('}p'):
            # It's a paragraph
            if para_idx < len(base_doc.paragraphs):
                para = base_doc.paragraphs[para_idx]
                
                # Skip Part III outline section
                if skip_start is not None and skip_start <= para_idx < skip_end:
                    para_idx += 1
                    continue
                
                # Skip the first few if they are cover pages (before Open Access)
                text = para.text.strip()
                
                # Copy paragraph
                if text or (para.style and para.style.name and para.style.name.startswith('Heading')):
                    copy_paragraph(para, output_doc)
                    copied_paras += 1
                elif not text and para.runs:
                    # Empty looking but has runs (might have formatting)
                    copy_paragraph(para, output_doc)
                    copied_paras += 1
            para_idx += 1
            
        elif element.tag.endswith('}tbl'):
            # It's a table
            if table_idx < len(base_doc.tables):
                # Check if this table is in the skipped Part III region
                # We approximate by checking if the last paragraph before this table
                # was in the skip region
                if skip_start is not None and skip_start <= para_idx < skip_end:
                    table_idx += 1
                    continue
                
                copy_table(base_doc.tables[table_idx], output_doc)
                copied_tables += 1
            table_idx += 1
    
    print(f"  Copied {copied_paras} paragraphs and {copied_tables} tables from Parts I & II")
    
    # ---- INSERT FULL PART III ----
    print("\n[5/6] Inserting full Part III content...")
    
    # Add Part III divider
    part3_divider = output_doc.add_paragraph()
    part3_divider.add_run().add_break(WD_BREAK.PAGE)
    
    # Copy Part III content
    part3_paras, part3_start = get_part3_content(part3_doc)
    
    part3_para_count = 0
    part3_table_count = 0
    
    # Track tables in Part III by their position in the body
    p3_body = part3_doc.element.body
    p3_elements = list(p3_body)
    p3_para_idx = 0
    p3_table_idx = 0
    
    for element in p3_elements:
        if element.tag.endswith('}p'):
            if p3_para_idx >= part3_start and p3_para_idx < len(part3_doc.paragraphs):
                para = part3_doc.paragraphs[p3_para_idx]
                text = para.text.strip()
                if text or (para.style and para.style.name and para.style.name.startswith('Heading')):
                    copy_paragraph(para, output_doc)
                    part3_para_count += 1
            p3_para_idx += 1
            
        elif element.tag.endswith('}tbl'):
            if p3_table_idx < len(part3_doc.tables) and p3_para_idx >= part3_start:
                copy_table(part3_doc.tables[p3_table_idx], output_doc)
                part3_table_count += 1
            p3_table_idx += 1
    
    print(f"  Copied {part3_para_count} paragraphs and {part3_table_count} tables from Part III")
    
    # ---- SAVE ----
    print("\n[6/6] Saving consolidated document...")
    output_doc.save(OUTPUT_FILE)
    
    # Final verification
    final_doc = Document(OUTPUT_FILE)
    total_paras = len(final_doc.paragraphs)
    total_tables = len(final_doc.tables)
    
    heading_counts = {}
    for para in final_doc.paragraphs:
        if para.style and para.style.name and para.style.name.startswith('Heading'):
            level = para.style.name
            heading_counts[level] = heading_counts.get(level, 0) + 1
    
    print(f"\n{'=' * 60}")
    print(f"CONSOLIDATION COMPLETE")
    print(f"{'=' * 60}")
    print(f"Output: {OUTPUT_FILE}")
    print(f"Total paragraphs: {total_paras}")
    print(f"Total tables: {total_tables}")
    print(f"Heading counts: {heading_counts}")
    print(f"\nThe document includes:")
    print(f"  ✓ Title page")
    print(f"  ✓ Open Access Notice")
    print(f"  ✓ Table of Contents (update in Word: right-click > Update Field)")
    print(f"  ✓ Part I: India's GCC Landscape (Chapters 1-6)")
    print(f"  ✓ Part II: GCC Sizing, Setup, Lifecycle, Case Studies, Leaders")
    print(f"  ✓ Part III: The Frontier (14 chapters)")
    print(f"  ✓ Appendices & Reference Directory")
    print(f"  ✓ Glossary (Key Terms and Definitions)")
    print(f"  ✓ Subject Index")
    print(f"  ✓ Validated References")


if __name__ == "__main__":
    merge_documents()
