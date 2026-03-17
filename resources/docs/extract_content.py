"""Extract structure and content overview from all docx files - output to file."""
import docx
import os

docs_dir = r"c:\Users\HP\OneDrive\Documents\GitHub\GCCLeadership\resources\docs"
output_path = r"c:\Users\HP\OneDrive\Documents\GitHub\GCCLeadership\resources\docs\analysis_output.txt"

files = [
    "GCC_Playbook_v_0.9.docx",
    "GCC_Playbook_Complete_CLEAN.docx", 
    "GCC_Playbook_2026_2030_Complete.docx",
    "GCC_Playbook_Part_III.docx",
    "GCC_Playbook_Part_II_CLEAN.docx",
]

with open(output_path, 'w', encoding='utf-8') as out:
    for fname in files:
        fpath = os.path.join(docs_dir, fname)
        out.write(f"\n{'='*80}\n")
        out.write(f"FILE: {fname}\n")
        out.write(f"{'='*80}\n")
        
        try:
            doc = docx.Document(fpath)
            
            props = doc.core_properties
            out.write(f"Title: {props.title}\n")
            out.write(f"Author: {props.author}\n")
            out.write(f"Total paragraphs: {len(doc.paragraphs)}\n")
            out.write(f"Total tables: {len(doc.tables)}\n")
            
            out.write(f"\n--- HEADING STRUCTURE ---\n")
            heading_count = 0
            for para in doc.paragraphs:
                if para.style and para.style.name and para.style.name.startswith('Heading'):
                    text = para.text.strip()
                    if text:
                        out.write(f"  {para.style.name}: {text[:150]}\n")
                        heading_count += 1
            out.write(f"Total headings: {heading_count}\n")
            
            # Print first non-empty paragraphs
            out.write(f"\n--- FIRST 20 NON-EMPTY PARAGRAPHS ---\n")
            count = 0
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    style = para.style.name if para.style else "None"
                    out.write(f"  [{style}] {text[:200]}\n")
                    count += 1
                    if count >= 20:
                        break
                        
        except Exception as e:
            out.write(f"Error: {e}\n")

print(f"Analysis written to {output_path}")
