"""
GCC Playbook PWA Content Extraction — v4
=========================================
Master: GCC_Playbook_v_1.0.docx  → Part I (uses List Paragraph + proper formatting)
Part II: GCC_Playbook_Part_II_CLEAN.docx → scale/setup/lifecycle/leadership + glossary
Part III: GCC_Playbook_Part_III.docx → 14 frontier chapters

Key improvement over v3:
- Handles 'List Paragraph' style → stored as bullet list arrays
- Handles run-level bold/italic → stored as inline markup hints
- Produces richer content[] entries that the PWA can format properly
"""
import json
import os
from docx import Document

DOCS_DIR   = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(os.path.dirname(DOCS_DIR), "..", "pwa", "public", "data")
os.makedirs(OUTPUT_DIR, exist_ok=True)

MASTER_FILE   = os.path.join(DOCS_DIR, "GCC_Playbook_v_1.0.docx")
PART2_FILE    = os.path.join(DOCS_DIR, "GCC_Playbook_Part_II_CLEAN.docx")
PART3_FILE    = os.path.join(DOCS_DIR, "GCC_Playbook_Part_III.docx")
COMPLETE_FILE = os.path.join(DOCS_DIR, "GCC_Playbook_Complete_CLEAN.docx")  # best glossary source

WML = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def para_type(para):
    """Return the logical type of a paragraph."""
    sn = para.style.name if para.style else "Normal"
    if sn.startswith("Heading 1"):  return "h1"
    if sn.startswith("Heading 2"):  return "h2"
    if sn.startswith("Heading 3"):  return "h3"
    if sn == "List Paragraph":       return "li"
    return "p"


def para_text(para):
    """Return stripped text of a paragraph."""
    return para.text.strip()


def is_appendix(title):
    return (title.startswith("A.") or "Appendix" in title
            or "Quick Navigation" in title or "How to Use" in title)


def is_key_section(title):
    return any(kw in title for kw in (
        "Chapter Summary", "Ten Things", "Ten Truths",
        "Ten Community", "Ten KPI", "Key Lessons",
        "Things Every", "Principles",
    ))


# ─────────────────────────────────────────────
# RICH CONTENT BLOCK
# ─────────────────────────────────────────────
def make_block(kind, text, items=None):
    """
    kind: 'p' | 'ul' | 'h3'
    For 'ul', items is a list of strings.
    For 'p'/'h3', text is the string.
    """
    if kind == "ul":
        return {"type": "ul", "items": items or []}
    return {"type": kind, "text": text}


# ─────────────────────────────────────────────
# CHAPTER EXTRACTION
# ─────────────────────────────────────────────
def extract_chapters(doc, prefix=""):
    chapters    = []
    cur_ch      = None
    cur_sec     = None
    cur_sub     = None
    pre_h1_secs = []
    seen_h1     = False

    # We collect consecutive 'li' paragraphs into a single ul block
    pending_list = []   # accumulate bullet items

    def flush_list(target):
        if pending_list and target is not None:
            block = make_block("ul", "", list(pending_list))
            if len(target.get("blocks", [])) < 30:
                target.setdefault("blocks", []).append(block)
        pending_list.clear()

    def add_p(target, text):
        if target is not None and len(target.get("blocks", [])) < 30:
            target.setdefault("blocks", []).append(make_block("p", text))

    def new_section(title):
        return {
            "title": title,
            "isKeyTakeaway": is_key_section(title),
            "blocks": [],
            "subsections": [],
        }

    def flush_pre(ch):
        if pre_h1_secs:
            ch["sections"].extend(pre_h1_secs)
            pre_h1_secs.clear()

    for para in doc.paragraphs:
        text = para_text(para)
        pt   = para_type(para)

        # ---- Heading 1 ----
        if pt == "h1":
            # flush any pending list into current target
            target = cur_sub or cur_sec or cur_ch
            flush_list(target)
            if cur_ch:
                chapters.append(cur_ch)
            seen_h1 = True
            cur_ch  = {"id": f"{prefix}ch-{len(chapters)+1}", "title": text,
                        "sections": [], "blocks": []}
            flush_pre(cur_ch)
            cur_sec = cur_sub = None

        # ---- Heading 2 ----
        elif pt == "h2":
            target = cur_sub or cur_sec or cur_ch
            flush_list(target)
            sec = new_section(text)
            if not seen_h1:
                pre_h1_secs.append(sec)
            elif cur_ch:
                cur_ch["sections"].append(sec)
            cur_sec = sec
            cur_sub = None

        # ---- Heading 3 ----
        elif pt == "h3":
            target = cur_sub or cur_sec or cur_ch
            flush_list(target)
            if not text:
                continue
            sub = {"title": text, "blocks": []}
            if cur_sec:
                cur_sec["subsections"].append(sub)
                cur_sub = sub
            elif cur_ch:
                sec = new_section(text)
                cur_ch["sections"].append(sec)
                cur_sec = sec
                cur_sub = None

        # ---- List Paragraph ----
        elif pt == "li":
            if text:
                pending_list.append(text)

        # ---- Body paragraph ----
        else:
            target = cur_sub or cur_sec or cur_ch
            # flush any accumulated list first
            if pending_list:
                flush_list(target)
            if text and len(text) >= 4:
                add_p(target, text)

    # flush tail
    target = cur_sub or cur_sec or cur_ch
    flush_list(target)
    if cur_ch:
        chapters.append(cur_ch)

    return chapters


# ─────────────────────────────────────────────
# GLOSSARY EXTRACTION
# ─────────────────────────────────────────────
def extract_glossary(doc):
    glossary    = []
    in_glossary = False
    paras       = []

    for para in doc.paragraphs:
        text  = para_text(para)
        style = para.style.name if para.style else "Normal"
        if "Key Terms and Definitions" in text:
            in_glossary = True
            continue
        if in_glossary:
            if "Subject Index" in text or ("Heading 1" in style and len(paras) > 5):
                break
            if style.startswith("Heading"):
                continue
            paras.append(text)

    # Strategy 1: inline separator
    for text in paras:
        if not text or len(text) < 10:
            continue
        for sep in [": ", " — ", " – ", " - "]:
            if sep in text:
                parts = text.split(sep, 1)
                term, defn = parts[0].strip(), parts[1].strip()
                if 3 < len(term) < 80 and len(defn) > 15:
                    glossary.append({"term": term, "definition": defn})
                    break

    # Strategy 2: pair pattern (short → long)
    if len(glossary) < 5:
        glossary = []
        i = 0
        while i < len(paras) - 1:
            term = paras[i].strip()
            defn = paras[i + 1].strip() if i + 1 < len(paras) else ""
            if term and defn and 3 < len(term) < 80 and len(defn) > 20 and len(term) < len(defn):
                glossary.append({"term": term, "definition": defn})
                i += 2
            else:
                i += 1

    return glossary


# ─────────────────────────────────────────────
# REFERENCES
# ─────────────────────────────────────────────
def extract_references(doc):
    refs, in_refs, cat = [], False, ""
    for para in doc.paragraphs:
        text  = para_text(para)
        style = para.style.name if para.style else "Normal"
        if any(kw in text for kw in ("VALIDATED REFERENCES", "Primary Sources")):
            in_refs = True; continue
        if in_refs:
            if not text: continue
            if style.startswith("Heading"): cat = text; continue
            refs.append({"category": cat, "text": text})
    return refs


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
def main():
    print("=" * 60)
    print("GCC Playbook PWA — Extraction v4 (v1.0 master)")
    print("=" * 60)

    print("\n[1/3] Part I  ← GCC_Playbook_v_1.0.docx")
    doc1   = Document(MASTER_FILE)
    all_p1 = extract_chapters(doc1, "p1-")
    p1_refs = extract_references(doc1)
    main_p1 = [c for c in all_p1 if not is_appendix(c["title"])]
    apx_p1  = [c for c in all_p1 if is_appendix(c["title"])]
    print(f"  {len(main_p1)} main chapters, {len(apx_p1)} appendix items")

    print("\n[2/3] Part II ← GCC_Playbook_Part_II_CLEAN.docx")
    doc2    = Document(PART2_FILE)
    main_p2 = extract_chapters(doc2, "p2-")
    p2_glos = extract_glossary(doc2)
    p2_refs = extract_references(doc2)
    print(f"  {len(main_p2)} chapters, {len(p2_glos)} glossary terms")

    print("\n[3/3] Part III ← GCC_Playbook_Part_III.docx")
    doc3    = Document(PART3_FILE)
    all_p3  = extract_chapters(doc3, "p3-")
    p3_refs = extract_references(doc3)
    main_p3 = [c for c in all_p3
               if "APPENDIX" not in c["title"].upper()
               and "VALIDATED REFERENCES" not in c["title"].upper()]
    print(f"  {len(main_p3)} chapters")

    # ── Glossary: use Complete_CLEAN which has the best formatted inline entries ──
    print("\n[+] Glossary ← GCC_Playbook_Complete_CLEAN.docx")
    doc_complete = Document(COMPLETE_FILE)
    glossary     = extract_glossary(doc_complete)
    if len(glossary) < 10:
        print(f"  Only {len(glossary)} terms from Complete_CLEAN, trying Part II...")
        glossary = extract_glossary(doc2)
    glossary = sorted(glossary, key=lambda x: x["term"].lower())
    print(f"  {len(glossary)} glossary terms extracted")

    all_refs, seen = [], set()
    for r in p1_refs + p2_refs + p3_refs:
        k = r["text"][:60]
        if k not in seen and r["text"].strip():
            seen.add(k); all_refs.append(r["text"])

    data = {
        "title":    "The Complete India GCC Reference 2026–2030",
        "subtitle": "Landscape | Maturity | Operating Models | AI | Deep Tech | M&A | Finance | ESG | 2030",
        "author":   "Curated by Kalilur Rahman",
        "parts": {
            "part1": {
                "title":    "Part I: India's GCC Landscape",
                "subtitle": "Landscape, Maturity, Operating Models, Challenges, KPIs, Community & Ecosystem",
                "chapters": main_p1,
            },
            "part2": {
                "title":    "Part II: GCC by Scale, Stage & Leadership",
                "subtitle": "Sizing, Setup, Lifecycle, Operating Models, Case Studies, Who's Who in GCC Leadership",
                "chapters": main_p2,
            },
            "part3": {
                "title":    "Part III: The Frontier",
                "subtitle": "AI Infrastructure, Deep Tech, M&A, Finance, Legal, Innovation, ESG, Career, 2030 Scenarios, Wellbeing",
                "chapters": main_p3,
            },
            "appendices": {
                "title":    "Appendices",
                "subtitle": "Decision Matrices, Scorecards, Compliance Calendar, KPI Benchmarks, Templates",
                "chapters": apx_p1,
            },
        },
        "glossary":   glossary,
        "references": all_refs[:150],
        "stats": {
            "totalChapters": len(main_p1) + len(main_p2) + len(main_p3),
            "part1Chapters": len(main_p1),
            "part2Chapters": len(main_p2),
            "part3Chapters": len(main_p3),
            "appendices":    len(apx_p1),
            "glossaryTerms": len(glossary),
            "references":    len(all_refs),
        },
    }

    out = os.path.join(OUTPUT_DIR, "content.json")
    with open(out, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    print(f"\n{'='*60}\nEXTRACTION COMPLETE\n{'='*60}")
    print(f"Output: {out}")
    print(f"Stats:  {json.dumps(data['stats'], indent=2)}")
    print(f"\nGlossary sample (first 5):")
    for g in glossary[:5]:
        print(f"  {g['term']}: {g['definition'][:70]}")

    # Verify block types in Part I chapter 1
    if main_p1:
        ch = main_p1[0]
        print(f"\nPart I ch1: '{ch['title']}'")
        print(f"  Top-level blocks: {len(ch.get('blocks',[]))}")
        for sec in ch.get('sections', [])[:3]:
            print(f"  Section '{sec['title'][:50]}': {len(sec.get('blocks',[]))} blocks, {len(sec.get('subsections',[]))} subsecs")
            for b in sec.get('blocks', [])[:3]:
                if b['type'] == 'ul':
                    print(f"    [ul] {len(b['items'])} items: {b['items'][0][:60] if b['items'] else ''}")
                else:
                    print(f"    [{b['type']}] {b['text'][:80]}")


if __name__ == "__main__":
    main()
