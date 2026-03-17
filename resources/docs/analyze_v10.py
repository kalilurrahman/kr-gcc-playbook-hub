"""Deep inspect v1.0 - look at heading structure + list paragraph samples."""
from docx import Document

doc = Document(r'c:\Users\HP\OneDrive\Documents\GitHub\GCCLeadership\resources\docs\GCC_Playbook_v_1.0.docx')

print("=== HEADINGS IN v1.0 ===")
heading_count = 0
for para in doc.paragraphs:
    text = para.text.strip()
    sn = para.style.name if para.style else 'Normal'
    if sn.startswith('Heading'):
        print(f"  {sn}: {text[:80]}")
        heading_count += 1
        if heading_count > 60:
            print("  ... (truncated)")
            break

print(f"\nTotal headings shown: {heading_count}")

print("\n=== SAMPLE - first Heading 1 and its content ===")
found_h1 = False
count = 0
for para in doc.paragraphs:
    text = para.text.strip()
    sn = para.style.name if para.style else 'Normal'
    if not found_h1 and sn == 'Heading 1':
        found_h1 = True
    if found_h1:
        fmt = ""
        if para.runs:
            r = para.runs[0]
            fmt = f"bold={r.bold} italic={r.italic}"
        print(f"  [{sn}] {fmt} | {repr(text[:100])}")
        count += 1
        if count > 30:
            break

print("\n=== LIST PARAGRAPH SAMPLES ===")
lp_count = 0
for para in doc.paragraphs:
    text = para.text.strip()
    sn = para.style.name if para.style else 'Normal'
    if sn == 'List Paragraph' and text:
        # Check numbering/indent
        indent = para.paragraph_format.left_indent
        numPr = para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
        has_num = numPr is not None
        print(f"  num={has_num} indent={indent} | {repr(text[:100])}")
        lp_count += 1
        if lp_count > 20:
            break
